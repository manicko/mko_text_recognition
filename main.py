from docx import Document
from pathlib import Path
from PIL import Image, ImageOps, ImageFilter, ImageEnhance
import pytesseract
import numpy as np
import cv2


# from processing.utils import spellcheck_with_hunspell


def deskew_image3(image):
    """Выравнивание изображения, если текст под углом."""
    # Определение угла поворота с использованием гистограммы проекций
    width, height = image.size
    data = image.getdata()
    angles = range(-10, 11)  # Проверяем углы от -10 до 10 градусов
    scores = []

    for angle in angles:
        rotated = image.rotate(angle, resample=Image.BICUBIC, expand=True)
        projection = [sum(rotated.crop((0, y, width, y + 1)).getdata()) for y in range(rotated.height)]
        scores.append((sum(projection), angle))

    _, best_angle = max(scores)
    return image.rotate(best_angle, resample=Image.BICUBIC, expand=True)




def process_image3(image_path, output_path):
    """Общая функция обработки изображения."""

    try:
    # Загрузка изображения
        image = Image.open(image_path)
        image = ImageOps.exif_transpose(image)  # Correct orientation using EXIF data
        image = image.convert("L")    # Приведение в градации серого
        image =ImageOps.autocontrast(image) # Нормализация яркости
        image = ImageEnhance.Brightness(image).enhance(1.05)  # Slightly increase brightness
        image = ImageEnhance.Contrast(image).enhance(1.05) # Усиление контраста
        image = image.filter(ImageFilter.MedianFilter(size=1)) # Удаление шума
        image = ImageEnhance.Sharpness(image).enhance(1.2)  # Sharpen
        # image = image.point(lambda x: 255 if x > 128 else 0, mode="1") # Бинаризация
        image = deskew_image3(image) # Выравнивание текста

        # Сохранение результата

        image_path = Path(output_path) / f"{Path(image_path).stem}_processed{Path(image_path).suffix}"
        image.save(image_path)
        print(f"Обработанное изображение сохранено: {image_path}")

        return image

    except Exception as e:
        print(f"Ошибка при обработке изображения {image_path}: {e}")
        return None




# cdf
def convert_to_grayscale(image):
    """Приведение изображения в градации серого."""
    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    return gray


def normalize_image(image):
    """Нормализация яркости изображения."""
    normalized = cv2.normalize(image, None, 0, 255, cv2.NORM_MINMAX)
    return normalized


def remove_noise(image):
    """Удаление шума с использованием фильтрации."""
    denoised = cv2.medianBlur(image, 5)
    return denoised


def rotate_image(image, angle):
    (h, w) = image.shape[: 2]
    center = (w // 2, h // 2)
    M = cv2.getRotationMatrix2D(center, angle, 1.0)
    corrected = cv2.warpAffine(image, M, (w, h), flags = cv2.INTER_CUBIC, \
        borderMode = cv2.BORDER_REPLICATE)
    return corrected

def determine_score(arr):
     histogram = np.sum(arr, axis = 2, dtype = float)
     score = np.sum((histogram[..., 1 :] - histogram[..., : -1]) ** 2, \
        axis = 1, dtype = float)
     return score

def deskew_image(image, delta = 0.1, limit = 5):
     thresh = cv2.threshold(image, 0, 255, cv2.THRESH_BINARY_INV + \
        cv2.THRESH_OTSU)[1]
     angles = np.arange(-limit, limit + delta, delta)
     img_stack = np.stack([rotate_image(thresh, angle) for angle \
        in angles], axis = 0)
     scores = determine_score(img_stack)
     best_angle = angles[np.argmax(scores)]
     corrected = rotate_image(image, best_angle)
     return corrected


def binarize_image(image):
    """Бинаризация изображения для выделения текста."""
    _, binary = cv2.threshold(image, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
    return binary


def enhance_contrast(image):
    """Усиление контраста с использованием CLAHE."""
    clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
    enhanced = clahe.apply(image)
    return enhanced


def process_image(image_path, output_path):
    """Общая функция обработки изображения."""
    try:
        # Загрузка изображения
        image = cv2.imread(image_path)

        # Приведение в градации серого
        image = convert_to_grayscale(image)

        # Нормализация яркости
        image = normalize_image(image)

        # Усиление контраста
        image = enhance_contrast(image)

        # Удаление шума
        image = remove_noise(image)

        # Бинаризация
        image = binarize_image(image)

        # Выравнивание текста
        image = deskew_image(image)

        # Сохранение результата
        result_image = Image.fromarray(image)
        processed_image_path = Path(output_path) / f"{Path(image_path).stem}_processed{Path(image_path).suffix}"
        result_image.save(processed_image_path)

        print(f"Обработанное изображение сохранено: {processed_image_path}")

        return result_image

    except Exception as e:
        print(f"Ошибка при обработке изображения {image_path}: {e}")
        raise e
    return None



def translate_files(path_in='data/in', path_out='data/out', lang_from="srp", ext='jpeg', save_as_one=False):
    """Обрабатывает изображения из папки и сохраняет текст и обработанные изображения."""
    path_in = Path(path_in)
    path_out = Path(path_out)
    path_out.mkdir(parents=True, exist_ok=True)  # Создать выходную папку, если её нет

    files = path_in.glob(f'*.{ext}')  # Получаем список файлов с указанным расширением
    combined_doc = Document() if save_as_one else None
    for file in files:
        print(f'Обработка файла {file}')

        # Обработка изображения
        processed_image = process_image3(file, path_out)
        if processed_image is None:
            continue
        try:
            # Распознавание текста с Tesseract
            # config = (f'--psm 6 '
            #           f'-c tessedit_char_whitelist=abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ0123456789@-.,:;čćžšđČĆŽŠĐ '
            #           f'-c preserve_interword_spaces=1')
            text = pytesseract.image_to_string(processed_image, lang=lang_from)

            # corrected_text = spellcheck_with_hunspell(text, lang_from=lang_from)


            if save_as_one:
                combined_doc.add_paragraph(f"File: {file.name}\n")
                combined_doc.add_paragraph(text)
                combined_doc.add_paragraph("\n---\n")
            else:
                doc = Document()
                doc.add_paragraph(text)
                doc.save(path_out / f"{file.stem}.docx")
                print(f"Текст успешно сохранён: {file.stem}.docx")
        except Exception as e:
            # print(f"Ошибка при обработке файла {file}: {e}")
            raise e

        if save_as_one:
            combined_doc.save(path_out / "combined_document.docx")
            print("Все тексты сохранены в общий файл: combined_document.docx")



if __name__ == '__main__':
    translate_files(lang_from='srp_latn', save_as_one=True)
